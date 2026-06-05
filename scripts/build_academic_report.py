from pathlib import Path
import textwrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as PdfImage, PageBreak


ROOT = Path(__file__).resolve().parents[1]
DOCX_OUT = ROOT / "ATIVIDADE_EXTENSIONISTA_II_TRABALHO_FINAL_LEANDRO_ANDALAFT.docx"
PDF_OUT = ROOT / "ATIVIDADE_EXTENSIONISTA_II_TRABALHO_FINAL_LEANDRO_ANDALAFT.pdf"
DIAGRAM_DIR = ROOT / "docs" / "diagramas"
EVIDENCE_DIR = ROOT / "docs" / "evidencias"


TITLE = "Atividade Extensionista II - Trabalho Final"
SUBTITLE = "Prontuario Unificado MDM"
STUDENT = "Leandro Andalaft dos Santos Junior"
RU = "4548358"


SECTIONS = [
    (
        "Identificacao",
        [
            "Curso: CST em Analise e Desenvolvimento de Sistemas.",
            "Disciplina: Atividade Extensionista II: Tecnologia Aplicada a Inclusao Digital - Projeto.",
            "Etapa: Trabalho final.",
            f"Aluno: {STUDENT}.",
            f"RU: {RU}.",
        ],
    ),
    (
        "Titulo",
        [
            "Desenvolvimento de uma ferramenta web responsiva para triagem, censo social, fila de atendimento e registro multiprofissional em acoes da Medicos do Mundo.",
        ],
    ),
    (
        "Setor de aplicacao",
        [
            "O projeto foi aplicado na Medicos do Mundo, organizacao da area da saude que realiza acoes humanitarias voltadas a populacao em situacao de vulnerabilidade social.",
            "O primeiro ciclo de uso ocorreu em Santos-SP, especialmente na regiao do Centro POP, bairro Paqueta. A partir da validacao em campo e da satisfacao dos voluntarios com a organizacao do fluxo, o sistema foi expandido para Itanhaem-SP, com localidade Belas Artes cadastrada como local de acao.",
            "A instituicao atua com frente multiprofissional de cuidado, incluindo cadastro, triagem, medicina, enfermagem e curativos, odontologia, farmacia, psicologia, fisioterapia, nutricao, vacinacao, biomedicina, veterinaria, podologia, justica de rua, acolhimento social, doacoes, apoio a mulher, emissao de documentos, beleza de rua e atendimento infantil/brinquedoteca.",
        ],
    ),
    (
        "Objetivos de Desenvolvimento Sustentavel",
        [
            "ODS 03 - Saude e bem-estar.",
            "ODS 10 - Reducao das desigualdades.",
            "ODS 17 - Parcerias e meios de implementacao.",
        ],
    ),
    (
        "Problema identificado",
        [
            "Antes da solucao, os registros das acoes eram fragmentados entre anotacoes, planilhas, mensagens e consolidacoes posteriores. Em ambiente de campo, isso dificultava acompanhar o fluxo de assistidos em tempo real, identificar pendencias, priorizar casos criticos e consolidar indicadores por local da acao.",
            "A ausencia de uma ferramenta unica tambem aumentava o retrabalho da equipe, pois o mesmo assistido poderia ter informacoes repetidas ou incompletas entre triagem, censo social e atendimentos por especialidade.",
        ],
    ),
    (
        "Objetivos",
        [
            "Mapear o fluxo atual de acolhimento, cadastro, triagem e atendimento.",
            "Identificar gargalos operacionais no registro, consulta, fila e consolidacao de dados.",
            "Definir requisitos funcionais para uso por voluntarios, academicos, profissionais, coordenacao e administracao.",
            "Desenvolver uma ferramenta simples, responsiva e segura para uso em celular ou notebook.",
            "Implantar um prototipo funcional com cadastro, triagem, censo, fila, atendimento, dashboard e exportacao.",
            "Validar o sistema em Santos-SP e expandir para Itanhaem-SP.",
        ],
    ),
    (
        "Metodologia",
        [
            "A metodologia adotada foi aplicada, iterativa e incremental. O desenvolvimento foi conduzido por ciclos de levantamento de requisitos, prototipacao funcional, validacao com usuarios, ajuste de interface, testes no navegador e refinamento de seguranca.",
            "O levantamento considerou comentarios da equipe multidisciplinar, prints de validacao, conversas de alinhamento, simulacoes com assistidos de teste e uso em campo. A cada ciclo, os formularios, alertas, textos, dashboards e fluxos de permissao foram ajustados para tornar o sistema mais intuitivo e adequado ao contexto de uma acao humanitaria.",
            "A validacao ocorreu por meio de cadastros, triagens, censos, atendimentos completos e parciais, atendimento extra, dashboard por local da acao e exportacao de dados.",
        ],
    ),
    (
        "Tecnologias utilizadas",
        [
            "Frontend: React, Vite, Tailwind CSS e Lucide React.",
            "Backend: Firebase Authentication, Cloud Firestore, Firebase Storage e Cloud Functions.",
            "Seguranca: Firestore Security Rules, validacao de e-mail, perfis de acesso e separacao por local da acao.",
            "Exportacao: biblioteca write-excel-file para planilhas com dashboards e abas por area.",
            "Publicacao e metricas: Vercel, Vercel Analytics e Vercel Speed Insights.",
            "Qualidade: ESLint, build Vite e testes funcionais no navegador.",
        ],
    ),
    (
        "Funcionalidades implementadas",
        [
            "Login com perfis operacionais e acesso protegido.",
            "Selecao obrigatoria do local da acao antes do uso operacional.",
            "Cadastro de assistidos com identificacao, nome civil/social, sexo ao nascer, genero e dados sociais.",
            "Triagem com queixa principal, uso de medicacao, sinais vitais, IMC, prioridade, atencao inclusiva e encaminhamentos.",
            "Censo social com moradia, rede de apoio, trabalho, uso de substancias, saude sexual/reprodutiva, antecedentes clinicos, saude mental, seguranca alimentar e pets.",
            "Fila de atendimento com indicacao de aguardando acao, finalizados e prioridade.",
            "Atendimento extra para situacoes que nao devem aguardar o fluxo completo.",
            "Prontuario com historico e registros por area.",
            "Dashboards gerais e por area, com exportacao geral e por data para coordenacao/administracao.",
        ],
    ),
    (
        "Resultados esperados e obtidos",
        [
            "O sistema foi aplicado inicialmente em Santos-SP e permitiu organizar melhor o fluxo de assistidos durante as acoes. A equipe passou a visualizar registros de triagem, censo social, atendimentos e pendencias de forma mais clara.",
            "Com a aceitacao em Santos-SP, o projeto foi expandido para Itanhaem-SP. A solucao passou a contemplar multiplos locais de acao, mantendo indicadores e prontuarios vinculados ao local onde o atendimento ocorreu.",
            "Como resultados, foram obtidos centralizacao das informacoes, reducao de retrabalho, apoio a priorizacao de casos, dashboards operacionais, exportacao estruturada e documentacao academica com evidencias sanitizadas.",
        ],
    ),
    (
        "LGPD e seguranca",
        [
            "O repositorio nao publica dados pessoais de assistidos ou voluntarios. As evidencias foram sanitizadas antes de serem incluidas na documentacao.",
            "O sistema possui autenticacao, regras de seguranca, perfis de acesso e cuidado para que exportacoes sensiveis sejam restritas a coordenacao e administracao.",
        ],
    ),
    (
        "Publicacao e execucao",
        [
            "Versao publicada: https://app-medicos.vercel.app/.",
            "Repositorio: https://github.com/Andalaftjr/app-medicos-mundo.",
            "Execucao local: instalar dependencias com npm install, criar arquivo .env a partir de .env.example e executar npm run dev.",
            "Validacao: npm run lint e npm run build.",
        ],
    ),
    (
        "Consideracoes finais",
        [
            "A atividade extensionista resultou em uma solucao funcional, aplicada e documentada, com impacto direto na organizacao operacional das acoes da Medicos do Mundo.",
            "O projeto demonstra a aplicacao pratica de tecnologia para inclusao digital e apoio a saude, conectando desenvolvimento de software, seguranca de dados, experiencia do usuario e responsabilidade social.",
        ],
    ),
]


EVIDENCES = [
    ("Dashboard operacional", "dashboard-operacional-sanitizado.png"),
    ("Indicadores operacionais", "indicadores-operacionais-sanitizado.png"),
    ("Fluxo de assistidos", "fluxo-assistidos-sanitizado.png"),
    ("Ficha do assistido", "ficha-assistido-sanitizado.png"),
    ("Triagem", "triagem-sanitizada.png"),
    ("Priorizacao e atencao inclusiva", "triagem-prioridade-sanitizada.png"),
    ("Censo social", "censo-social-sanitizado.png"),
    ("Uso de substancias no censo", "censo-substancias-sanitizado.png"),
]

ACCENTS = [
    ("Prontuario", "Prontuário"),
    ("Medicos", "Médicos"),
    ("Analise", "Análise"),
    ("Inclusao", "Inclusão"),
    ("aplicacao", "aplicação"),
    ("aplicacoes", "aplicações"),
    ("Aplicacao", "Aplicação"),
    ("acao", "ação"),
    ("acoes", "ações"),
    ("Acao", "Ação"),
    ("Itanhaem", "Itanhaém"),
    ("Paqueta", "Paquetá"),
    ("regiao", "região"),
    ("saude", "saúde"),
    ("Saude", "Saúde"),
    ("sao", "são"),
    ("Sao", "São"),
    ("nao", "não"),
    ("Nao", "Não"),
    ("area", "área"),
    ("areas", "áreas"),
    ("Area", "Área"),
    ("populacao", "população"),
    ("solucao", "solução"),
    ("Solucao", "Solução"),
    ("atencao", "atenção"),
    ("Atencao", "Atenção"),
    ("priorizacao", "priorização"),
    ("edicao", "edição"),
    ("execucao", "execução"),
    ("Execucao", "Execução"),
    ("codigo", "código"),
    ("Repositorio", "Repositório"),
    ("repositorio", "repositório"),
    ("autenticacao", "autenticação"),
    ("Autenticacao", "Autenticação"),
    ("selecao", "seleção"),
    ("Selecao", "Seleção"),
    ("obrigatoria", "obrigatória"),
    ("identificacao", "identificação"),
    ("Identificacao", "Identificação"),
    ("genero", "gênero"),
    ("modulos", "módulos"),
    ("Modulos", "Módulos"),
    ("tambem", "também"),
    ("prontuarios", "prontuários"),
    ("Prontuarios", "Prontuários"),
    ("pendencias", "pendências"),
    ("possui", "possui"),
    ("reducao", "redução"),
    ("Reducao", "Redução"),
    ("implementacao", "implementação"),
    ("Implementacao", "Implementação"),
    ("organizacao", "organização"),
    ("informacoes", "informações"),
    ("consolidacao", "consolidação"),
    ("coordenacao", "coordenação"),
    ("administracao", "administração"),
    ("triagem", "triagem"),
    ("historico", "histórico"),
    ("Historico", "Histórico"),
    ("clinico", "clínico"),
    ("clinicas", "clínicas"),
    ("tecnica", "técnica"),
    ("Tecnica", "Técnica"),
    ("tecnico", "técnico"),
    ("Tecnologias", "Tecnologias"),
    ("seguranca", "segurança"),
    ("Seguranca", "Segurança"),
    ("evidencias", "evidências"),
    ("Evidencias", "Evidências"),
    ("validacao", "validação"),
    ("Validacao", "Validação"),
    ("necessarias", "necessárias"),
    ("necessario", "necessário"),
    ("permissoes", "permissões"),
    ("exportacao", "exportação"),
    ("Exportacao", "Exportação"),
    ("publicacao", "publicação"),
    ("Publicacao", "Publicação"),
    ("responsavel", "responsável"),
    ("usuarios", "usuários"),
    ("voluntarios", "voluntários"),
    ("academicos", "acadêmicos"),
    ("profissionais", "profissionais"),
    ("operacao", "operação"),
    ("decisao", "decisão"),
    ("vinculacao", "vinculação"),
    ("multiplos", "múltiplos"),
    ("unica", "única"),
    ("proprio", "próprio"),
    ("prototipacao", "prototipação"),
    ("iterativa", "iterativa"),
    ("incremental", "incremental"),
    ("assistent", "assistent"),
]


def human(text):
    for src, dst in ACCENTS:
        text = text.replace(src, dst)
    return text


def load_font(size=22, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            pass
    return ImageFont.load_default()


def draw_wrapped(draw, xy, text, font, fill, max_width, line_spacing=6):
    text = human(text)
    words = text.split()
    lines = []
    current = ""
    for word in words:
        tentative = f"{current} {word}".strip()
        if draw.textbbox((0, 0), tentative, font=font)[2] <= max_width:
            current = tentative
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_spacing
    return y


def draw_box(draw, xy, text, fill, outline, font, text_fill="#0b2545"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    text_w = x2 - x1 - 28
    y = y1 + 22
    draw_wrapped(draw, (x1 + 14, y), human(text), font, text_fill, text_w)


def arrow(draw, start, end, fill="#26315f"):
    draw.line([start, end], fill=fill, width=4)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        points = [(x2, y2), (x2 - 14, y2 - 8), (x2 - 14, y2 + 8)]
    else:
        points = [(x2, y2), (x2 + 14, y2 - 8), (x2 + 14, y2 + 8)]
    draw.polygon(points, fill=fill)


def make_diagrams():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    font_title = load_font(30, True)
    font_box = load_font(21, True)
    font_small = load_font(18, False)
    bg = "#f8fafc"
    navy = "#26315f"
    blue = "#e7efff"
    green = "#e8fff3"
    amber = "#fff7df"

    diagrams = []

    img = Image.new("RGB", (1400, 760), bg)
    d = ImageDraw.Draw(img)
    d.text((50, 40), human("Diagrama de contexto"), font=font_title, fill=navy)
    boxes = [
        ((70, 160, 330, 285), "Assistido / paciente", amber),
        ((70, 420, 330, 545), "Equipe da acao", amber),
        ((540, 285, 860, 425), "Sistema web responsivo MDM", blue),
        ((1060, 160, 1320, 285), "Firebase Auth, Firestore e Functions", green),
        ((1060, 420, 1320, 545), "Dashboards e exportacao Excel", green),
    ]
    for xy, text, color in boxes:
        draw_box(d, xy, text, color, "#cbd5e1", font_box)
    arrow(d, (330, 222), (540, 335))
    arrow(d, (330, 482), (540, 375))
    arrow(d, (860, 335), (1060, 222))
    arrow(d, (860, 375), (1060, 482))
    draw_wrapped(d, (540, 490), "Voluntarios, academicos, profissionais, coordenacao e administracao acessam conforme perfil.", font_small, "#334155", 540)
    path = DIAGRAM_DIR / "diagrama-contexto.png"
    img.save(path)
    diagrams.append(path)

    img = Image.new("RGB", (1500, 900), bg)
    d = ImageDraw.Draw(img)
    d.text((50, 35), human("Fluxo operacional da acao"), font=font_title, fill=navy)
    steps = [
        ("Login", 80, 150),
        ("Selecionar local da acao", 360, 150),
        ("Cadastrar ou localizar assistido", 720, 150),
        ("Triagem e sinais vitais", 1080, 150),
        ("Censo social e historico", 80, 440),
        ("Fila e encaminhamento", 420, 440),
        ("Atendimento por area", 760, 440),
        ("Dashboard e exportacao", 1100, 440),
    ]
    for label, x, y in steps:
        draw_box(d, (x, y, x + 270, y + 125), label, blue if y == 150 else green, "#cbd5e1", font_box)
    for i in range(3):
        arrow(d, (steps[i][1] + 270, steps[i][2] + 62), (steps[i + 1][1], steps[i + 1][2] + 62))
    arrow(d, (1215, 275), (215, 440))
    for i in range(4, 7):
        arrow(d, (steps[i][1] + 270, steps[i][2] + 62), (steps[i + 1][1], steps[i + 1][2] + 62))
    draw_wrapped(d, (80, 735), "Regra central: indicadores e prontuarios sao vinculados ao local da acao, nao apenas a filial do voluntario.", font_small, "#334155", 1320)
    path = DIAGRAM_DIR / "fluxo-operacional.png"
    img.save(path)
    diagrams.append(path)

    img = Image.new("RGB", (1500, 860), bg)
    d = ImageDraw.Draw(img)
    d.text((50, 35), human("Arquitetura tecnica e seguranca"), font=font_title, fill=navy)
    architecture = [
        ((70, 170, 360, 300), "Celular / notebook", amber),
        ((510, 170, 800, 300), "React + Vite + Tailwind", blue),
        ((950, 80, 1260, 210), "Firebase Authentication", green),
        ((950, 270, 1260, 400), "Cloud Firestore + Rules", green),
        ((950, 460, 1260, 590), "Cloud Functions", green),
        ((510, 520, 800, 650), "Excel, dashboards e evidencias", blue),
    ]
    for xy, text, color in architecture:
        draw_box(d, xy, text, color, "#cbd5e1", font_box)
    arrow(d, (360, 235), (510, 235))
    arrow(d, (800, 210), (950, 145))
    arrow(d, (800, 250), (950, 335))
    arrow(d, (800, 290), (950, 525))
    arrow(d, (950, 525), (800, 585))
    draw_wrapped(d, (70, 720), "RBAC, validacao de e-mail, filtros por localAcao e restricao de exportacao reduzem risco de acesso indevido.", font_small, "#334155", 1350)
    path = DIAGRAM_DIR / "arquitetura-seguranca.png"
    img.save(path)
    diagrams.append(path)
    return diagrams


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(human(text))
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(human(item))


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, value in enumerate(rows[0]):
        set_cell_text(hdr[i], value, True)
        set_cell_shading(hdr[i], "E8EEF5")
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def fit_image_width(path, max_width_inches=6.2):
    with Image.open(path) as img:
        w, h = img.size
    ratio = min(max_width_inches, 6.2)
    return Inches(ratio)


def build_docx(diagrams):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in [("Heading 1", 15, "26315F"), ("Heading 2", 12.5, "1D4ED8"), ("Heading 3", 11.5, "334155")]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(human(TITLE))
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("26315F")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(human(SUBTITLE))
    sr.bold = True
    sr.font.size = Pt(15)
    sr.font.color.rgb = RGBColor.from_string("1D4ED8")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(human(f"{STUDENT} | RU {RU}\nCST em Analise e Desenvolvimento de Sistemas\nMedicos do Mundo - Santos-SP e Itanhaem-SP"))
    doc.add_paragraph()

    add_table(
        doc,
        [
            ["Campo", "Descricao"],
            ["Disciplina", "Atividade Extensionista II: Tecnologia Aplicada a Inclusao Digital - Projeto"],
            ["Etapa", "Trabalho final"],
            ["ODS", "03 Saude e bem-estar; 10 Reducao das desigualdades; 17 Parcerias"],
            ["Publicacao", "https://app-medicos.vercel.app/"],
            ["Repositorio", "https://github.com/Andalaftjr/app-medicos-mundo"],
        ],
    )

    for heading, paragraphs in SECTIONS:
        doc.add_heading(human(heading), level=1)
        if heading in {"Objetivos", "Tecnologias utilizadas", "Funcionalidades implementadas", "LGPD e seguranca", "Publicacao e execucao"}:
            add_bullets(doc, paragraphs)
        else:
            for paragraph in paragraphs:
                doc.add_paragraph(human(paragraph))

    doc.add_page_break()
    doc.add_heading(human("Diagramas do projeto"), level=1)
    diagram_captions = [
        "Diagrama de contexto: participantes, sistema, Firebase e exportacoes.",
        "Fluxo operacional: da entrada no app ate dashboards e exportacao.",
        "Arquitetura tecnica e seguranca: frontend, Firebase, regras e funcoes.",
    ]
    for caption, path in zip(diagram_captions, diagrams):
        doc.add_heading(human(caption), level=2)
        doc.add_picture(str(path), width=Inches(6.1))

    doc.add_page_break()
    doc.add_heading(human("Evidencias de desenvolvimento e aplicacao"), level=1)
    doc.add_paragraph(human("As evidencias abaixo foram sanitizadas para evitar exposicao de dados pessoais, documentos e informacoes clinicas identificaveis."))
    for caption, filename in EVIDENCES:
        path = EVIDENCE_DIR / filename
        if path.exists():
            doc.add_heading(human(caption), level=2)
            doc.add_picture(str(path), width=Inches(4.7))

    doc.core_properties.author = STUDENT
    doc.core_properties.title = TITLE
    doc.save(DOCX_OUT)


def pdf_paragraph(text):
    text = human(text)
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf(diagrams):
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        rightMargin=1.6 * cm,
        leftMargin=1.6 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.4 * cm,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CenterTitle", parent=styles["Title"], alignment=TA_CENTER, textColor=colors.HexColor("#26315F"), fontSize=20, leading=24, spaceAfter=8))
    styles.add(ParagraphStyle(name="Sub", parent=styles["Normal"], alignment=TA_CENTER, textColor=colors.HexColor("#334155"), fontSize=10, leading=13, spaceAfter=14))
    styles.add(ParagraphStyle(name="H1Custom", parent=styles["Heading1"], textColor=colors.HexColor("#26315F"), fontSize=14, leading=18, spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle(name="BodyCustom", parent=styles["BodyText"], fontSize=9.5, leading=13, spaceAfter=6))
    styles.add(ParagraphStyle(name="BulletCustom", parent=styles["BodyText"], fontSize=9.5, leading=13, leftIndent=14, bulletIndent=4, spaceAfter=4))

    story = [
        Paragraph(pdf_paragraph(TITLE), styles["CenterTitle"]),
        Paragraph(pdf_paragraph(SUBTITLE), styles["Sub"]),
        Paragraph(pdf_paragraph(f"{STUDENT} | RU {RU}") + "<br/>" + pdf_paragraph("CST em Analise e Desenvolvimento de Sistemas") + "<br/>" + pdf_paragraph("Medicos do Mundo - Santos-SP e Itanhaem-SP"), styles["Sub"]),
        Spacer(1, 8),
    ]
    meta_table = Table(
        [
            ["Campo", "Descricao"],
            ["Disciplina", human("Atividade Extensionista II: Tecnologia Aplicada a Inclusao Digital - Projeto")],
            ["Etapa", "Trabalho final"],
            ["ODS", "03 Saude e bem-estar; 10 Reducao das desigualdades; 17 Parcerias"],
            ["Publicacao", "https://app-medicos.vercel.app/"],
            ["Repositorio", "https://github.com/Andalaftjr/app-medicos-mundo"],
        ],
        colWidths=[3.2 * cm, 13 * cm],
    )
    meta_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#26315F")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.extend([meta_table, Spacer(1, 12)])

    for heading, paragraphs in SECTIONS:
        story.append(Paragraph(pdf_paragraph(heading), styles["H1Custom"]))
        if heading in {"Objetivos", "Tecnologias utilizadas", "Funcionalidades implementadas", "LGPD e seguranca", "Publicacao e execucao"}:
            for item in paragraphs:
                story.append(Paragraph(pdf_paragraph(item), styles["BulletCustom"], bulletText="-"))
        else:
            for paragraph in paragraphs:
                story.append(Paragraph(pdf_paragraph(paragraph), styles["BodyCustom"]))

    story.append(PageBreak())
    story.append(Paragraph(pdf_paragraph("Diagramas do projeto"), styles["H1Custom"]))
    captions = [
        "Diagrama de contexto: participantes, sistema, Firebase e exportacoes.",
        "Fluxo operacional: da entrada no app ate dashboards e exportacao.",
        "Arquitetura tecnica e seguranca: frontend, Firebase, regras e funcoes.",
    ]
    for caption, path in zip(captions, diagrams):
        story.append(Paragraph(pdf_paragraph(caption), styles["BodyCustom"]))
        story.append(PdfImage(str(path), width=16.5 * cm, height=9.0 * cm))
        story.append(Spacer(1, 8))

    story.append(PageBreak())
    story.append(Paragraph(pdf_paragraph("Evidencias de desenvolvimento e aplicacao"), styles["H1Custom"]))
    story.append(Paragraph(pdf_paragraph("As evidencias abaixo foram sanitizadas para evitar exposicao de dados pessoais, documentos e informacoes clinicas identificaveis."), styles["BodyCustom"]))
    for caption, filename in EVIDENCES:
        path = EVIDENCE_DIR / filename
        if path.exists():
            story.append(Paragraph(pdf_paragraph(caption), styles["BodyCustom"]))
            story.append(PdfImage(str(path), width=10.2 * cm, height=13.0 * cm))
            story.append(Spacer(1, 8))

    doc.build(story)


def main():
    diagrams = make_diagrams()
    build_docx(diagrams)
    build_pdf(diagrams)
    print(f"DOCX: {DOCX_OUT}")
    print(f"PDF: {PDF_OUT}")
    print("Diagramas:")
    for path in diagrams:
        print(f"- {path}")


if __name__ == "__main__":
    main()
